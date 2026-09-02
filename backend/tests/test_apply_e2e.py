"""End-to-end test of the apply loop through the real FastAPI app.

Covers everything except the network hop to the LLM: resume storage, job lookup, tailor,
guardrail, DOCX render, application logging, and the download endpoint. The provider is
stubbed so the test is deterministic and needs no API key — but every other layer is the
real one, including SQLite and python-docx.
"""

from __future__ import annotations

import docx

BASE_RESUME = {
    "basics": {"name": "Jane Doe", "email": "jane@example.com", "summary": "Engineer."},
    "work": [
        {
            "name": "Acme Corp",
            "position": "Software Engineer",
            "startDate": "2022-01",
            "endDate": "2024-06",
            "highlights": ["Reduced p95 latency by 15% with a Redis cache."],
        }
    ],
    "education": [],
    "skills": [{"name": "Languages", "level": "", "keywords": ["Python"]}],
    "projects": [],
}

JOB = {
    "id": "greenhouse:acme:999",
    "source": "greenhouse",
    "title": "Senior Backend Engineer",
    "company": "Globex",
    "location": "Remote",
    "remote": True,
    "apply_url": "https://example.com/apply/999",
    "description": "Python, Redis, and Kubernetes.",
}


class _Stub:
    """Returns whatever resume it was given — honest or fabricated, per the test.

    Deliberately does NOT subclass LLMProvider or import app.schemas at module scope.
    The `client` fixture purges `app.*` from sys.modules to get a clean instance per
    test, which means any class imported up here would be a *stale* copy that pydantic
    rejects as a different type. Building the model inside the call keeps us on whatever
    module generation the app is currently running.
    """

    name = "stub"

    def __init__(self, fabricate: bool = False) -> None:
        self.fabricate = fabricate

    async def complete_structured(self, *, system, user, schema, max_tokens=16000):
        resume = schema.model_validate(BASE_RESUME)
        resume.basics.summary = "Backend engineer specializing in Python and Redis."
        if self.fabricate:
            resume.skills.append(
                type(resume.skills[0])(name="Infra", keywords=["Kubernetes"])
            )
        return resume

    async def health(self):
        return True, "stub"


def _seed(client, monkeypatch, fabricate: bool = False) -> None:
    """Store a base resume, cache the job, and pin the LLM to the stub."""
    assert client.post(
        "/api/resumes", json={"name": "Base", "data": BASE_RESUME}
    ).status_code == 200

    from app.db import SessionLocal
    from app.models import Job

    with SessionLocal() as db:
        db.add(
            Job(
                id=JOB["id"],
                source=JOB["source"],
                title=JOB["title"],
                company=JOB["company"],
                location=JOB["location"],
                remote=JOB["remote"],
                apply_url=JOB["apply_url"],
                description=JOB["description"],
            )
        )
        db.commit()

    import app.routers.apply as apply_module

    monkeypatch.setattr(apply_module, "build_provider", lambda db: _Stub(fabricate))


def test_health(client) -> None:
    assert client.get("/api/health").json()["status"] == "ok"


def test_apply_without_resume_is_a_clear_error(client, monkeypatch) -> None:
    from app.db import SessionLocal
    from app.models import Job

    with SessionLocal() as db:
        db.add(Job(id=JOB["id"], source="x", title="T", company="C", apply_url="u"))
        db.commit()

    import app.routers.apply as apply_module

    monkeypatch.setattr(apply_module, "build_provider", lambda db: _Stub())
    res = client.post(f"/api/apply/{JOB['id']}")
    assert res.status_code == 400
    assert "base resume" in res.json()["detail"].lower()


def test_apply_produces_a_downloadable_docx(client, monkeypatch, tmp_path) -> None:
    _seed(client, monkeypatch)

    res = client.post(f"/api/apply/{JOB['id']}")
    assert res.status_code == 200, res.text
    body = res.json()

    assert body["tailoring"]["fell_back"] is False
    assert body["tailoring"]["violations"] == []
    assert body["docx_url"]

    download = client.get(body["docx_url"])
    assert download.status_code == 200
    assert download.headers["content-type"].startswith(
        "application/vnd.openxmlformats"
    )

    # Read the real file back the way an ATS would.
    out = tmp_path / "downloaded.docx"
    out.write_bytes(download.content)
    text = "\n".join(p.text for p in docx.Document(str(out)).paragraphs)
    assert "Jane Doe" in text
    assert "Acme Corp" in text
    assert "15%" in text


def test_apply_logs_the_application(client, monkeypatch) -> None:
    _seed(client, monkeypatch)
    client.post(f"/api/apply/{JOB['id']}")

    rows = client.get("/api/applications").json()
    assert len(rows) == 1
    assert rows[0]["company"] == "Globex"
    assert rows[0]["status"] == "prepared"
    assert rows[0]["apply_url"] == JOB["apply_url"]


def test_fabricating_model_falls_back_and_warns(client, monkeypatch, tmp_path) -> None:
    """The whole point of the system: a bad model must not produce a bad resume."""
    _seed(client, monkeypatch, fabricate=True)

    body = client.post(f"/api/apply/{JOB['id']}").json()

    assert body["tailoring"]["fell_back"] is True
    assert body["tailoring"]["warning"]
    assert any(v["value"] == "Kubernetes" for v in body["tailoring"]["violations"])

    # And the rendered file must be the honest one — no Kubernetes in the DOCX.
    out = tmp_path / "fallback.docx"
    out.write_bytes(client.get(body["docx_url"]).content)
    text = "\n".join(p.text for p in docx.Document(str(out)).paragraphs)
    assert "Kubernetes" not in text
    assert "Python" in text


def test_apply_produces_a_single_page_pdf(client, monkeypatch, tmp_path) -> None:
    """PDF now works on every install and must be exactly one page — the guarantee."""
    import pypdfium2 as pdfium

    _seed(client, monkeypatch)

    body = client.post(f"/api/apply/{JOB['id']}").json()
    assert body["pdf_url"]
    assert body["pdf_error"] is None

    download = client.get(body["pdf_url"])
    assert download.status_code == 200
    assert download.headers["content-type"] == "application/pdf"

    out = tmp_path / "downloaded.pdf"
    out.write_bytes(download.content)
    assert download.content[:5] == b"%PDF-"  # real PDF header, not an error page

    pdf = pdfium.PdfDocument(str(out))
    try:
        assert len(pdf) == 1, f"PDF must be a single page, got {len(pdf)}"
    finally:
        pdf.close()


def test_pdf_render_failure_is_reported_not_fatal(client, monkeypatch) -> None:
    """A render bug must degrade to DOCX-only for that apply, not 500 the request."""
    import app.services.render_pdf as pdf_module
    import app.routers.apply as apply_module

    _seed(client, monkeypatch)
    monkeypatch.setattr(
        apply_module,
        "render_pdf",
        lambda resume, path: (_ for _ in ()).throw(
            pdf_module.PDFRenderError("simulated render failure")
        ),
    )

    res = client.post(f"/api/apply/{JOB['id']}")
    assert res.status_code == 200
    body = res.json()
    assert body["docx_url"]
    assert body["pdf_url"] is None
    assert "simulated render failure" in body["pdf_error"]


def test_download_rejects_unknown_format(client, monkeypatch) -> None:
    _seed(client, monkeypatch)
    app_id = client.post(f"/api/apply/{JOB['id']}").json()["application_id"]
    assert client.get(f"/api/download/{app_id}/exe").status_code == 400


def test_settings_never_returns_the_api_key(client) -> None:
    client.put("/api/settings", json={"llm_provider": "openai", "api_key": "sk-secret"})
    body = client.get("/api/settings").json()
    assert body["has_key"] is True
    assert "sk-secret" not in str(body)


# ---------------------------------------------------------------- cover letters


class _CoverLetterStub:
    """Same "no module-level app.* imports" discipline as `_Stub`, for the same reason."""

    name = "stub"

    def __init__(self, fabricate: bool = False) -> None:
        self.fabricate = fabricate

    async def complete_structured(self, *, system, user, schema, max_tokens=16000):
        body = (
            "Dear Hiring Team, I bring hands-on Python experience from my time at Acme "
            "Corp. Sincerely, Jane Doe"
        )
        if self.fabricate:
            body = (
                "Dear Hiring Team, I improved performance by 90% at Acme Corp. "
                "Sincerely, Jane Doe"
            )
        return schema.model_validate({"body": body})

    async def health(self):
        return True, "stub"


def test_cover_letter_requires_existing_application(client) -> None:
    from app.db import SessionLocal
    from app.models import Job

    with SessionLocal() as db:
        db.add(Job(id=JOB["id"], source="x", title="T", company="C", apply_url="u"))
        db.commit()

    res = client.post(f"/api/apply/{JOB['id']}/cover-letter")
    assert res.status_code == 404


def test_cover_letter_endpoint_returns_body(client, monkeypatch) -> None:
    _seed(client, monkeypatch)
    client.post(f"/api/apply/{JOB['id']}")

    import app.routers.apply as apply_module

    monkeypatch.setattr(apply_module, "build_provider", lambda db: _CoverLetterStub())

    res = client.post(f"/api/apply/{JOB['id']}/cover-letter")
    assert res.status_code == 200, res.text
    body = res.json()
    assert body["body"]
    assert body["fell_back"] is False


def test_cover_letter_fabrication_falls_back(client, monkeypatch) -> None:
    _seed(client, monkeypatch)
    client.post(f"/api/apply/{JOB['id']}")

    import app.routers.apply as apply_module

    monkeypatch.setattr(
        apply_module, "build_provider", lambda db: _CoverLetterStub(fabricate=True)
    )

    res = client.post(f"/api/apply/{JOB['id']}/cover-letter")
    assert res.status_code == 200, res.text
    body = res.json()
    assert body["fell_back"] is True
    assert body["warning"]
    assert "90%" not in body["body"]


# ---------------------------------------------------------------------- ad-hoc jobs


def test_adhoc_job_is_idempotent(client) -> None:
    payload = {
        "company": "Initech",
        "title": "Backend Engineer",
        "description": "Build things.",
        "apply_url": "https://boards.example.com/initech/backend",
    }
    first = client.post("/api/jobs/adhoc", json=payload)
    assert first.status_code == 200, first.text
    assert first.json()["id"].startswith("adhoc:")
    assert first.json()["source"] == "adhoc"

    second = client.post("/api/jobs/adhoc", json=payload)
    assert second.status_code == 200
    assert second.json()["id"] == first.json()["id"], "resubmitting must upsert, not duplicate"

    rows = client.get("/api/applications").json()  # sanity: no crash from the extra row
    assert isinstance(rows, list)


# ------------------------------------------------------------------ applications by-url


def test_applications_by_url_round_trip(client, monkeypatch) -> None:
    _seed(client, monkeypatch)
    client.post(f"/api/apply/{JOB['id']}")

    import app.routers.apply as apply_module

    monkeypatch.setattr(apply_module, "build_provider", lambda db: _CoverLetterStub())
    client.post(f"/api/apply/{JOB['id']}/cover-letter")

    # Query-string and fragment noise must not break the match.
    res = client.get(
        "/api/applications/by-url", params={"url": JOB["apply_url"] + "?utm=1#top"}
    )
    assert res.status_code == 200, res.text
    body = res.json()
    assert body["application"]["job_id"] == JOB["id"]
    assert body["resume"]["basics"]["name"] == "Jane Doe"
    assert body["cover_letter"]


def test_applications_by_url_404_when_unknown(client) -> None:
    res = client.get("/api/applications/by-url", params={"url": "https://example.com/nope"})
    assert res.status_code == 404
